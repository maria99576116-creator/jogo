import json
from docx import Document
from docx.shared import Inches

def generate_word_doc(json_file_path, docx_file_path):
    document = Document()
    document.add_heading('Jogo de Cartas: Perguntas e Respostas sobre Literatura Brasileira', level=1)

    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    for level, questions in data.items():
        document.add_heading(f'Nível: {level.capitalize()}', level=2)
        for i, q in enumerate(questions):
            document.add_heading(f'{i+1}. Pergunta: {q["pergunta"]}', level=3)
            document.add_paragraph(f'Resposta Correta: {q["resposta"]}')
            document.add_paragraph('Opções:')
            for option in q['opcoes']:
                document.add_paragraph(f'- {option}', style='List Bullet')
            document.add_paragraph('\n') # Add a blank line for spacing

    document.save(docx_file_path)

if __name__ == '__main__':
    json_path = '/home/ubuntu/jogo-literatura-brasileira/src/data/perguntas.json'
    docx_path = 'perguntas_literatura_brasileira.docx'
    generate_word_doc(json_path, docx_path)

